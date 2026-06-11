import asyncio
import re
from .exceptions import UnsupportedMimeTypeError, EmptyDocumentError, ExtractionError

class DocumentExtractor:
    """
    Extract plain text from PDF, DOCX, TXT files.
    Returns clean text ready for token counting + LLM context.
    """

    @staticmethod
    async def extract(file_bytes: bytes, mime_type: str, filename: str) -> str:
        """Route to correct extractor based on MIME type"""
        if mime_type == 'application/pdf':
            return await DocumentExtractor._extract_pdf(file_bytes)
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return await DocumentExtractor._extract_docx(file_bytes)
        elif mime_type == 'text/plain':
            return await DocumentExtractor._extract_txt(file_bytes)
        else:
            raise UnsupportedMimeTypeError(f"Unsupported type: {mime_type}")

    @staticmethod
    async def _extract_pdf(file_bytes: bytes) -> str:
        """pdfplumber — handles text PDFs and basic layout"""
        import pdfplumber
        import io
        loop = asyncio.get_event_loop()

        def _sync_extract():
            text_parts = []
            try:
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text(x_tolerance=3, y_tolerance=3)
                        if text:
                            text_parts.append(text.strip())
            except Exception as e:
                raise ExtractionError(f"PDF extraction failed: {str(e)}")
            return '\n\n'.join(text_parts)

        # Run in thread pool — pdfplumber is sync/blocking
        text = await loop.run_in_executor(None, _sync_extract)
        if not text.strip():
            raise EmptyDocumentError("PDF appears to be image-only or empty")
        return DocumentExtractor._clean_text(text)

    @staticmethod
    async def _extract_docx(file_bytes: bytes) -> str:
        """python-docx — extract paragraphs + tables"""
        import docx
        import io
        loop = asyncio.get_event_loop()

        def _sync_extract():
            try:
                doc = docx.Document(io.BytesIO(file_bytes))
                parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        parts.append(para.text.strip())
                # Also extract table text
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            parts.append(row_text)
                return '\n'.join(parts)
            except Exception as e:
                raise ExtractionError(f"DOCX extraction failed: {str(e)}")

        text = await loop.run_in_executor(None, _sync_extract)
        if not text.strip():
            raise EmptyDocumentError("DOCX appears empty")
        return DocumentExtractor._clean_text(text)

    @staticmethod
    async def _extract_txt(file_bytes: bytes) -> str:
        """TXT — detect encoding then decode"""
        import chardet
        detected = chardet.detect(file_bytes)
        encoding = detected.get('encoding') or 'utf-8'
        try:
            text = file_bytes.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            text = file_bytes.decode('utf-8', errors='replace')
        if not text.strip():
            raise EmptyDocumentError("Text file is empty")
        return DocumentExtractor._clean_text(text)

    @staticmethod
    def _clean_text(text: str) -> str:
        """Remove excessive whitespace, null bytes, control chars"""
        text = text.replace('\x00', '')  # null bytes
        text = re.sub(r'\r\n', '\n', text)  # normalize line endings
        text = re.sub(r'\n{3,}', '\n\n', text)  # max 2 consecutive newlines
        text = re.sub(r'[ \t]{2,}', ' ', text)  # multiple spaces → one
        text = re.sub(r'[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]', '', text)  # remove non-printable
        return text.strip()
